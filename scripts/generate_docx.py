"""
Generates NOVU_Builder_Prezentacni_Dokument.docx from the markdown source.
Run: python scripts/generate_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MD_PATH = Path(__file__).parent.parent / "docs" / "NOVU_Builder_Prezentacni_Dokument.md"
OUT_PATH = Path(__file__).parent.parent / "docs" / "NOVU_Builder_Prezentacni_Dokument.docx"

# ── colour palette ────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black navy
C_ACCENT = RGBColor(0x0F, 0x3F, 0x7F)   # corporate blue
C_LIGHT  = RGBColor(0xF4, 0xF6, 0xFA)   # table header background (simulated)
C_GREY   = RGBColor(0x55, 0x55, 0x55)   # body text / captions
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN  = RGBColor(0x1A, 0x73, 0x48)
C_RED    = RGBColor(0xB0, 0x00, 0x20)


def set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F3F7F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_styles(doc: Document) -> None:
    """Apply base style overrides."""
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = C_GREY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(14)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = C_DARK
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = C_ACCENT
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = C_DARK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)


def add_cover_page(doc: Document) -> None:
    """Insert a simple branded cover page."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NOVU Builder")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = C_ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Prezentační a obchodní dokument")
    run2.font.name = "Calibri"
    run2.font.size = Pt(16)
    run2.font.color.rgb = C_DARK

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    meta_lines = [
        ("Verze systému:", "v0.7.001"),
        ("Datum dokumentu:", "2. dubna 2026"),
        ("Určeno pro:", "Investory · Obchodní partnery · Pilotní zákazníky"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}  ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = C_DARK
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.color.rgb = C_GREY

    doc.add_page_break()


def apply_inline_formatting(run_text: str, paragraph) -> None:
    """Parse **bold**, `code`, and plain text, adding runs to paragraph."""
    # Split on **bold** and `code`
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = token_re.split(run_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = C_DARK
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                r = paragraph.add_run(part)
    # ensure all runs inherit normal size if not set
    for run in paragraph.runs:
        if not run.font.size:
            run.font.size = Pt(10.5)


def render_table(doc: Document, lines: list[str]) -> None:
    """Render a markdown pipe table as a styled Word table."""
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-| :]+\|\s*$", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    # Pad all rows
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            if i == 0:
                # Header row
                set_cell_bg(cell, "0F3F7F")
                p = cell.paragraphs[0]
                p.clear()
                apply_inline_formatting(cell_text, p)
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = C_WHITE
                    run.font.size = Pt(10)
            else:
                bg = "F4F6FA" if i % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.clear()
                apply_inline_formatting(cell_text, p)
                for run in p.runs:
                    if not run.font.color.rgb:
                        run.font.color.rgb = C_GREY

    # Set approximate column widths
    total_width = Inches(6.0)
    col_width = total_width / col_count
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    doc.add_paragraph()  # spacing after table


def parse_and_render(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    ordered_counter: dict[int, int] = {}

    while i < len(lines):
        line = lines[i]

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r"^---+\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Page break hint (we insert after each main section heading) ──────
        # handled below by detecting H2

        # ── Headings ─────────────────────────────────────────────────────────
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # ── Pipe table ───────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            render_table(doc, table_lines)
            continue

        # ── Ordered list ─────────────────────────────────────────────────────
        m_ol = re.match(r"^(\d+)\.\s+(.*)", line)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            apply_inline_formatting(m_ol.group(2), p)
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────────────
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            bullet_text = re.sub(r"^[-*]\s+", "", line)
            apply_inline_formatting(bullet_text, p)
            i += 1
            continue

        # ── Italic metadata footer lines ─────────────────────────────────────
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("*"))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        apply_inline_formatting(line, p)
        i += 1


def set_page_margins(doc: Document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NOVU Builder  ·  Prezentační dokument  ·  v0.7.001  ·  2. dubna 2026")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    set_page_margins(doc)
    configure_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    parse_and_render(doc, md_text)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
